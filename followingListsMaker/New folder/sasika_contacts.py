# IMPORT LIBRARIES
import datetime, time, random, os, csv
import pyautogui as pt
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.chrome.options import Options
chromeoptions = Options()
import docx
doc = docx.Document()

geckoDriverPath = os.getcwd()
##driver = webdriver.Firefox()
start = datetime.datetime.now()
driver = webdriver.Chrome(options=chromeoptions,executable_path=str(geckoDriverPath)+'\\chromedriver.exe')
driver.get('http://www.instagram.com/')

# CREDENTIALS
myEmail = 'sasikapamith2016@gmail.com'
myPassword = '.^ML!r{^Jk9/8(Uu'

username = WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.CSS_SELECTOR, "input[name='username']")))
password = WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.CSS_SELECTOR, "input[name='password']")))

username.clear()
password.clear()
username.send_keys(myEmail)
password.send_keys(myPassword)

log_in = WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.CSS_SELECTOR, "button[type='submit']"))).click()
print("Log in completed...")

# DISMISS THE NOTIFICATIONS
not_now = WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.XPATH, "//button[contains(text(),'Not Now')]"))).click()
not_now2 = WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.XPATH, "//button[contains(text(),'Not Now')]"))).click()
print('Dismissed Notifications...')

names = []
numbers = []
driver.get('https://www.instagram.com/accounts/contact_history/')
time.sleep(10)
notAdded = []
for i in range(3002):
    j = i+1
    time.sleep(1)
    try:
        name = driver.find_element_by_xpath("//div[1]/div[5]/div[%s]/div/div[1]/h1" % j).text
        number = driver.find_element_by_xpath("//div[1]/div[5]/div[%s]/div/div[2]/h1" % j).text
        names.append(name)
        numbers.append(number)
        print(j,name,number)
        doc.add_paragraph(str('%s' %j) +" " +  str(name) + " " + str(number))
        doc.add_paragraph(" ")
    except:
        notAdded.append(j)
        pass
doc.save('insta_sasika_amarasinghe_contacts.docx')

print('Total names = ',len(names))
print('Total contacts = ',len(numbers))
print('No of errors while executing = ',len(notAdded))
